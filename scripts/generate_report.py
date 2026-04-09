import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_revised_report():
    doc = Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'  # Professional modern font
    font.size = Pt(10.5)

    # Title Section
    title = doc.add_heading('Dual-Mode Legal AI Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Enterprise-Grade Legal RAG: Architecture, Orchestration, and Compliance Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    # 1. System Guarantee & Current Progress
    doc.add_heading('1. System Guarantee & Overview', level=1)
    doc.add_paragraph(
        "PROJECT GUARANTEE: \"No answer is returned without verifiable, grounded evidence from an isolated source.\"",
        style='Caption'
    )
    doc.add_paragraph(
        "Current development has successfully integrated a 13-node autonomous graph, ensuring "
        "that every legal query is validated, classified, and cross-verified against a strictly isolated provenance layer."
    )

    # 2. Advanced Retrieval Strategy (Gaps 2, 7, 8, 9)
    doc.add_heading('2. High-Precision Retrieval Infrastructure', level=1)
    doc.add_paragraph(
        "We have moved beyond basic vector search to a hybrid retrieval architecture designed for "
        "statutory precision and private document integrity."
    )
    doc.add_paragraph("- Hybrid Search: Integration of BM25 (keyword) and Vector (semantic) search for precise statute lookup.", style='List Bullet')
    doc.add_paragraph("- Reranking & Top-K Tuning: Initial retrieval of top-30 candidates, followed by post-retrieval reranking to the top-5 most relevant spans.", style='List Bullet')
    doc.add_paragraph("- Metadata-Level Isolation: Rigid filtering at the database layer using `title_number` (Federal) and `document_id` (Private), preventing cross-contamination.", style='List Bullet')
    doc.add_paragraph("- Structure-Aware Chunking: 750-token window with 10% overlap and 'Hierarchy Preservation' logic (retaining headings and page numbers in every chunk).", style='List Bullet')
    doc.add_paragraph("- Incremental Indexing: Periodic automated ingestion runs ensure statutory updates (e.g., Title 11 changes) are reflected without full system rebuilds.", style='List Bullet')

    # 3. Agent Orchestration & State Design (Gaps 10, 11, 12)
    doc.add_heading('3. The 13-Node Orchestration Graph', level=1)
    doc.add_paragraph(
        "The system uses LangGraph to manage a durable GraphState. This state object persists decisions, "
        "retrieval plans, and validation issues, allowing nodes to pass data-driven control flow."
    )
    nodes = [
        "Ingest: Query normalization and intent detection.",
        "Classify Mode: Strict routing between Federal (U.S. Code) and Document contexts.",
        "Extract Entities: Capturing § sections and statutory citations.",
        "Mode Hints: Narrowing search via Title-level filtering (e.g., Title 18 for crimes).",
        "Document Scope: Pinpointing specific pages/articles in uploaded files.",
        "Make Plan: Crafting the hybrid search and reranking parameters.",
        "Route Tool: Identity node for tool-level security enforcement.",
        "Retrieve: Execution of isolated hybrid search in Qdrant.",
        "Grade: Evaluating average relevance score against a strict 0.7 threshold.",
        "Generate: Grounded response using GPT-4o with span-level citations.",
        "Verify: Independent verification node using GPT-4o-mini to audit the draft.",
        "Retry Loop: Logic-driven retry node (explained below).",
        "Finalize & Persist: Telemetry storage in PostgreSQL/Prometheus."
    ]
    for n in nodes:
        doc.add_paragraph(n, style='List Number')

    # 4. Failure Handling & Retry Policy (Gaps 1, 11)
    doc.add_heading('4. Failure Handling & Resiliency', level=1)
    doc.add_paragraph("- Failure Path: In the event of empty retrieval, low relevance (< 0.65), or verification failure, the system returns an explicit 'Insufficient Evidence' response.", style='List Bullet')
    doc.add_paragraph("- Retry Policy: One retry loop (at Node 12) is permitted. Between retries, the agent broadens search parameters (e.g., increasing top-k or relaxing title filters) to locate missing context.", style='List Bullet')
    doc.add_paragraph("- Logged Failures: Every 'Insufficient Evidence' result includes a specific failure reason (e.g., 'Verification Hallucination Detected') in the technical telemetry.", style='List Bullet')

    # 5. Deep Verification Accountability (Gap 3)
    doc.add_heading('5. The Verification Audit (LLM-as-Judge)', level=1)
    doc.add_paragraph(
        "Verification isn't a vague check; it is a multi-step audit of the 'Draft Answer' against the 'Raw Evidence':"
    )
    doc.add_paragraph("- Span Comparison: The verifier ensures the draft doesn't exceed the semantic boundaries of the retrieved text.", style='List Bullet')
    doc.add_paragraph("- Citation Mapping: Every Material Claim is mapped 1:1 to a citation. If a claim lacks support, it is flagged.", style='List Bullet')
    doc.add_paragraph("- Mode Enforcement: Federal statues are rejected in Document mode to prevent general knowledge hallucination.", style='List Bullet')

    # 6. Evaluation, Performance & Cost (Gaps 4, 5, 6)
    doc.add_heading('6. Performance, Cost, and Benchmarking', level=1)
    doc.add_paragraph("- Accuracy Metrics: Tracking 'Grounding Success Rate' and 'Cite-per-Answer' density via Prometheus.", style='List Bullet')
    doc.add_paragraph("- Latency Targets: Target response time of sub-5 seconds. Parallel ingestion optimized for 2MB/s processing via Celery.", style='List Bullet')
    doc.add_paragraph("- Cost Optimization: Usage of GPT-4o-mini for repetitive verification nodes reduces token costs by ~80% compared to a full GPT-4 pipeline.", style='List Bullet')

    # 7. Production Infrastructure & Security (Gaps 13, 14, 15)
    doc.add_heading('7. Production Readiness & Security', level=1)
    doc.add_paragraph("- Infrastructure: Scalable Docker Compose deployment, optimized for Oracle Cloud ARM instances (24GB RAM).", style='List Bullet')
    doc.add_paragraph("- Multi-Tenancy: Hard isolation of `upload_id` at the Qdrant collection layer ensures user data never crosses private sessions.", style='List Bullet')
    doc.add_paragraph("- Rate Limiting: Redis-based rate limiting prevents abuse and ensures high availability.", style='List Bullet')

    # 8. UX, Feedback & Narrative (Gaps 16, 17, 18, 19, 20)
    doc.add_heading('8. Experience & Long-term Strategy', level=1)
    doc.add_paragraph("- Explainability UI: Frontend includes a 'Citation Drawer' that highlights the exact text source in real-time.", style='List Bullet')
    doc.add_paragraph("- Human-in-the-Loop: Integrated thumbs up/down feedback that stores corrections for future model fine-tuning.", style='List Bullet')
    doc.add_paragraph("- Baseline Comparison: Unlike 'Generalist LLMs' or 'Basic RAG,' our system prevents context mixing and enforces legal compliance via audit trails.", style='List Bullet')

    doc.add_paragraph('\n')
    doc.add_paragraph('Confidential - Revised for Professional Implementation Delivery', style='Caption').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(os.path.dirname(script_dir), 'Legal_AI_Agent_Report_Revised.docx')
    doc.save(output_path)
    print(f"Successfully generated revised report at: {output_path}")

if __name__ == "__main__":
    create_revised_report()
